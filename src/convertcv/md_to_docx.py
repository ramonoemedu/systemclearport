#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the markdown file located next to this script
base_dir = os.path.dirname(os.path.abspath(__file__))
md_file = os.path.join(base_dir, 'convertpdftotext.md')
with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Parse markdown and add to document
lines = content.split('\n')
for line in lines:
    line = line.rstrip()
    
    if not line:
        # Add empty paragraph for spacing
        doc.add_paragraph()
    elif line.startswith('# '):
        # H1 - Main title
        heading = doc.add_heading(line[2:], level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    elif line.startswith('## '):
        # H2 - Section heading
        heading = doc.add_heading(line[3:], level=1)
        for run in heading.runs:
            run.font.size = Pt(12)
            run.font.bold = True
    elif line.startswith('### '):
        # H3 - Subsection
        heading = doc.add_heading(line[4:], level=2)
        for run in heading.runs:
            run.font.size = Pt(11)
            run.font.bold = True
    elif line.startswith('- '):
        # Bullet point
        paragraph = doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('+ '):
        # Plus bullet point (treat same as dash)
        paragraph = doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('* '):
        # Asterisk bullet point
        paragraph = doc.add_paragraph(line[2:], style='List Bullet')
    else:
        # Regular paragraph
        doc.add_paragraph(line)

# Save the document next to the script
output_file = os.path.join(base_dir, 'convertpdftotext.docx')
doc.save(output_file)
print(f"Word document created successfully: {output_file}")

