#!/usr/bin/env python3
import os
import sys
from html import escape

base_dir = os.path.dirname(os.path.abspath(__file__))
pdf_path = os.path.join(base_dir, 'Ramon_OEM_CV_2026.pdf')
output_path = os.path.join(base_dir, 'Ramon_OEM_CV_2026.docx')

def extract_with_pymupdf(path):
    try:
        import fitz
    except Exception:
        return None
    doc = fitz.open(path)
    pages = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text('text')
        pages.append(text)
    return pages

def extract_with_pdfminer(path):
    try:
        from pdfminer.high_level import extract_text
    except Exception:
        return None
    try:
        text = extract_text(path)
        # pdfminer returns whole text; split heuristically by form-feed or pages
        pages = text.split('\f') if '\f' in text else [text]
        return pages
    except Exception:
        return None

def write_docx_from_pages(pages, out_path):
    try:
        from docx import Document
    except Exception as e:
        print('python-docx missing:', e)
        return False
    doc = Document()
    for idx, page_text in enumerate(pages, start=1):
        doc.add_heading(f'Page {idx}', level=1)
        # split into paragraphs
        for para in [p.strip() for p in page_text.split('\n\n') if p.strip()]:
            doc.add_paragraph(para)
    doc.save(out_path)
    return True

def main():
    if not os.path.exists(pdf_path):
        print('PDF not found at', pdf_path)
        sys.exit(1)

    pages = extract_with_pymupdf(pdf_path)
    if pages is None:
        pages = extract_with_pdfminer(pdf_path)

    if pages is None:
        print('No PDF text extractor available (fitz or pdfminer).')
        sys.exit(2)

    ok = write_docx_from_pages(pages, output_path)
    if not ok:
        print('Failed to write docx. Ensure python-docx is installed.')
        sys.exit(3)

    print('Created Word document at', output_path)

if __name__ == '__main__':
    main()
